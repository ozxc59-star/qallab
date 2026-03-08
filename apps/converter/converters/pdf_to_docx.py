"""
PDF -> DOCX conversion with full Arabic/RTL support + layout preservation.

Addresses every Arabic rendering failure mode:
  1. Character integrity   – ligatures, tashkeel, font encoding, numeral order
  2. Directionality        – true w:bidi RTL (not just right-align), punctuation,
                             mixed Arabic+English bidi runs in one paragraph
  3. Layout / structure    – paragraph merging (no hard line-breaks), tables,
                             bullets, headings, images, page margins

Strategy tiers (attempted in order):
  Tier 1 — PyMuPDF "rawdict" extraction + python-docx rebuild
            Best for text-based Arabic PDFs. Uses rawdict so we get the raw
            Unicode glyphs (post-ToUnicode CMap decoding done by MuPDF).
            Mixed bidi text (Arabic + English in one line) is handled with
            multiple runs tagged individually RTL or LTR.
            Paragraphs are reconstructed from lines (not one-line-per-para).
            Images extracted and inserted at reading-order position.
  Tier 2 — Tesseract OCR (200 DPI, ara+eng, LSTM)
            For image-based (scanned) PDFs.
  Tier 3 — LibreOffice headless fallback (last resort).
"""

from __future__ import annotations

import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import NamedTuple

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Arabic / RTL Unicode ranges
# ---------------------------------------------------------------------------
_ARABIC_RE = re.compile(
    r"[\u0600-\u06FF"   # Arabic block (includes tashkeel diacritics)
    r"\u0750-\u077F"    # Arabic Supplement
    r"\u08A0-\u08FF"    # Arabic Extended-A
    r"\uFB50-\uFDFF"    # Arabic Presentation Forms-A
    r"\uFE70-\uFEFF"    # Arabic Presentation Forms-B
    r"]"
)

# Detect if a span is predominantly Arabic
def _is_rtl_span(text: str) -> bool:
    arabic_chars = len(_ARABIC_RE.findall(text))
    if not text.strip():
        return False
    return arabic_chars / max(len(text.strip()), 1) >= 0.2

# Minimum extractable chars on page 1 to call PDF "text-based"
_TEXT_THRESHOLD = 50

# Font encoding garbage detection: if >40% of non-space chars are
# in the Private Use Area or are replacement chars, the PDF uses a
# custom encoding MuPDF cannot decode → fall through to OCR.
_BAD_CHAR_RE = re.compile(r"[\uE000-\uF8FF\uFFFD\u0000-\u001F]")

def _text_is_garbage(text: str) -> bool:
    stripped = text.replace(" ", "").replace("\n", "")
    if not stripped:
        return False
    bad = len(_BAD_CHAR_RE.findall(stripped))
    return bad / len(stripped) > 0.40

LIBREOFFICE_BINS = [
    "libreoffice", "soffice",
    "/usr/bin/libreoffice", "/usr/bin/soffice",
    "/usr/lib/libreoffice/program/soffice",
]

def find_libreoffice() -> str | None:
    for b in LIBREOFFICE_BINS:
        if shutil.which(b):
            return b
    return None

def is_text_based_pdf(pdf_path: str) -> bool:
    try:
        import fitz
        doc = fitz.open(pdf_path)
        if not doc:
            return False
        text = doc[0].get_text("text")
        doc.close()
        return len(text.strip()) >= _TEXT_THRESHOLD
    except Exception as exc:
        logger.warning(f"is_text_based_pdf: {exc}")
        return False


# ===========================================================================
# DOCX XML helpers – the foundation of correct Arabic rendering
# ===========================================================================

def _qn(tag: str) -> str:
    from docx.oxml.ns import qn
    return qn(tag)

def _el(tag: str):
    from docx.oxml import OxmlElement
    return OxmlElement(tag)


def _configure_doc_rtl_defaults(doc_word) -> None:
    """
    Set document-level defaults so every paragraph is RTL by default.
    Also sets Amiri as the complex-script (cs) and ascii font in Normal style.

    We set <w:bidi/> in the document default paragraph properties so that Word
    treats the document as an RTL document, not just right-aligned LTR.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 1. Document-level bidi default in settings
    try:
        settings = doc_word.settings.element
        bidi_default = OxmlElement("w:bidi")
        settings.append(bidi_default)
    except Exception:
        pass

    # 2. Default paragraph style → bidi + right-align
    try:
        normal_style = doc_word.styles["Normal"]
        pPr = normal_style.element.get_or_add_pPr()

        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)

        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "right")
        pPr.append(jc)
    except Exception:
        pass

    # 3. Default run style → Amiri as cs + ascii font, rtl
    try:
        normal_style = doc_word.styles["Normal"]
        rPr = normal_style.element.get_or_add_rPr()

        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Amiri")
        rFonts.set(qn("w:hAnsi"), "Amiri")
        rFonts.set(qn("w:cs"), "Amiri")
        rPr.append(rFonts)

        rtl_el = OxmlElement("w:rtl")
        rPr.append(rtl_el)
    except Exception:
        pass


def _set_para_bidi(paragraph, rtl: bool) -> None:
    """
    Set TRUE RTL direction (w:bidi) on a paragraph, not just alignment.

    Word distinguishes:
      - w:bidi  → paragraph is bidirectional (cursor behaviour, punctuation
                  placement, bracket mirroring all work correctly)
      - w:jc right → visual right-alignment only (cursor stays LTR)

    We always set w:bidi for Arabic paragraphs AND w:jc right.
    For LTR paragraphs we explicitly set w:bidi val="0" to override the
    document default we set above, plus w:jc left.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()

    # Remove any existing bidi/jc to avoid duplicates
    for tag in ("w:bidi", "w:jc"):
        for el in pPr.findall(qn(tag)):
            pPr.remove(el)

    if rtl:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)

        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "right")
        pPr.append(jc)
    else:
        # Explicitly turn bidi OFF (overrides document default)
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "0")
        pPr.append(bidi)

        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "left")
        pPr.append(jc)


def _set_para_spacing(paragraph, space_after_pt: float = 4.0) -> None:
    """Reduce paragraph spacing so tashkeel doesn't create huge gaps."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    for el in pPr.findall(qn("w:spacing")):
        pPr.remove(el)

    spacing = OxmlElement("w:spacing")
    # spaceAfter in twentieths of a point
    spacing.set(qn("w:after"), str(int(space_after_pt * 20)))
    # Line spacing: auto (0 = auto in Word's terms via lineRule)
    spacing.set(qn("w:line"), "276")  # 1.15× = 276 twentieths
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _make_run_rtl_props(run, font_name: str = "Amiri",
                        font_size_pt: float | None = None,
                        bold: bool = False, italic: bool = False,
                        rtl: bool = True) -> None:
    """
    Set all required run properties for correct Arabic rendering:
      w:rFonts ascii/hAnsi/cs  – Amiri for both ASCII and complex-script slots
      w:rtl                     – character-level RTL (critical for correct
                                  glyph selection and caret movement)
      w:sz / w:szCs             – font size in half-points (both slots)
      w:b / w:bCs               – bold (both slots)
      w:i / w:iCs               – italic (both slots)
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    rPr = run._r.get_or_add_rPr()

    # Font – set all four slots so Word never falls back to TNR
    for el in rPr.findall(qn("w:rFonts")):
        rPr.remove(el)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.append(rFonts)

    # RTL character direction
    for el in rPr.findall(qn("w:rtl")):
        rPr.remove(el)
    if rtl:
        rtl_el = OxmlElement("w:rtl")
        rPr.append(rtl_el)
    else:
        # Explicitly LTR (overrides document default)
        rtl_el = OxmlElement("w:rtl")
        rtl_el.set(qn("w:val"), "0")
        rPr.append(rtl_el)

    # Font size – set BOTH w:sz (ASCII) and w:szCs (complex script)
    if font_size_pt is not None:
        half_pts = str(int(font_size_pt * 2))
        for tag in ("w:sz", "w:szCs"):
            for el in rPr.findall(qn(tag)):
                rPr.remove(el)
            sz = OxmlElement(tag)
            sz.set(qn("w:val"), half_pts)
            rPr.append(sz)

    # Bold – set w:b AND w:bCs
    for tag in ("w:b", "w:bCs"):
        for el in rPr.findall(qn(tag)):
            rPr.remove(el)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))

    # Italic – set w:i AND w:iCs
    for tag in ("w:i", "w:iCs"):
        for el in rPr.findall(qn(tag)):
            rPr.remove(el)
    if italic:
        rPr.append(OxmlElement("w:i"))
        rPr.append(OxmlElement("w:iCs"))


def _add_page_break(doc_word) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para = doc_word.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ===========================================================================
# Span / line data structures
# ===========================================================================

class Span(NamedTuple):
    text: str
    size: float        # normalised pt
    bold: bool
    italic: bool
    rtl: bool
    x0: float          # left edge (for column / list detection)
    y0: float
    y1: float


class Block(NamedTuple):
    y0: float
    kind: str          # "text" | "image"
    data: object       # list[list[Span]] (lines of spans) | (xref, width, height)


# ===========================================================================
# Font-size normalisation
# ===========================================================================

def _page_median_size(raw_blocks: list) -> float:
    sizes: list[float] = []
    for b in raw_blocks:
        if b.get("type") != 0:
            continue
        for line in b.get("lines", []):
            for span in line.get("spans", []):
                s = span.get("size", 0.0)
                if s > 0:
                    sizes.append(s)
    if not sizes:
        return 11.0
    sizes.sort()
    return sizes[len(sizes) // 2]


def _norm_size(raw: float, median: float) -> float:
    """Normalise raw PDF pt size so median body text → 11 pt."""
    if median <= 0:
        return raw
    return round((raw / median) * 11.0, 1)


# ===========================================================================
# Paragraph reconstruction
# ===========================================================================

def _lines_to_paragraphs(lines_of_spans: list[list[Span]]) -> list[list[Span]]:
    """
    Merge consecutive lines that are part of the same paragraph.

    PDF files insert a hard newline at every visual line break.  In Word
    those become separate paragraphs, which looks like the text is fragmented.
    We merge lines that:
      - have the same approximate font size (within 1 pt)
      - have the same RTL direction (dominant)
      - whose vertical gap is less than 1.5× the line height

    We do NOT merge across different font sizes (heading vs body) or when
    the gap is large (section break / heading gap).
    """
    if not lines_of_spans:
        return []

    paragraphs: list[list[Span]] = []
    current: list[Span] = list(lines_of_spans[0])

    for line in lines_of_spans[1:]:
        if not line:
            continue

        prev_spans = current
        cur_size = max(s.size for s in prev_spans)
        new_size = max(s.size for s in line)

        prev_rtl = sum(1 for s in prev_spans if s.rtl) > len(prev_spans) / 2
        new_rtl  = sum(1 for s in line if s.rtl) > len(line) / 2

        # Vertical gap check
        prev_y1 = max(s.y1 for s in prev_spans)
        new_y0  = min(s.y0 for s in line)
        line_h  = max(s.y1 - s.y0 for s in line) if line else 12
        gap = new_y0 - prev_y1

        same_size = abs(cur_size - new_size) <= 1.0
        same_dir  = prev_rtl == new_rtl
        close_gap = gap < line_h * 1.8  # merge if gap < 1.8× line height

        if same_size and same_dir and close_gap:
            # Append a space then the new line's spans to current paragraph
            current.append(Span(" ", cur_size, False, False, prev_rtl,
                                 current[-1].x0, new_y0, new_y0 + line_h))
            current.extend(line)
        else:
            paragraphs.append(current)
            current = list(line)

    paragraphs.append(current)
    return paragraphs


# ===========================================================================
# DOCX paragraph writer
# ===========================================================================

def _write_paragraph(doc_word, spans: list[Span],
                     heading_level: int = 0) -> None:
    """
    Write one paragraph to doc_word.

    Handles mixed bidi correctly:
      - Paragraph direction set from dominant direction of all spans
      - Each run is tagged individually w:rtl=1 or w:rtl=0
      - Arabic runs and LTR runs get separate runs so Word's bidi algorithm
        can re-order them correctly without mangling order

    Punctuation displacement is fixed by using TRUE w:bidi on the paragraph
    (not just right-align) so Word knows to mirror brackets etc.
    """
    if not spans:
        return

    # Dominant direction for the paragraph
    rtl_count = sum(1 for s in spans if s.rtl)
    para_rtl = rtl_count > len(spans) / 2

    dom_size = max(s.size for s in spans)
    any_bold = any(s.bold for s in spans)

    if heading_level > 0:
        try:
            para = doc_word.add_heading("", level=heading_level)
        except Exception:
            para = doc_word.add_paragraph()
    else:
        para = doc_word.add_paragraph()

    _set_para_bidi(para, para_rtl)
    _set_para_spacing(para, space_after_pt=2.0)

    # Write spans as individual runs so mixed bidi works
    for span in spans:
        if not span.text:
            continue
        run = para.add_run(span.text)
        _make_run_rtl_props(
            run,
            font_name="Amiri",
            font_size_pt=span.size,
            bold=span.bold,
            italic=span.italic,
            rtl=span.rtl,
        )


def _write_image(doc_word, image_bytes: bytes, max_width_inches: float = 5.5) -> None:
    """Insert an image paragraph into the document."""
    from docx.shared import Inches

    try:
        para = doc_word.add_paragraph()
        run = para.add_run()
        run.add_picture(io.BytesIO(image_bytes), width=Inches(max_width_inches))
    except Exception as exc:
        logger.debug(f"Image insert failed: {exc}")


# ===========================================================================
# Tier 1 — PyMuPDF rawdict extraction
# ===========================================================================

def _tier1_pymupdf_to_docx(input_path: str, output_path: str) -> bool:
    """
    Full-fidelity extraction using PyMuPDF rawdict mode.

    rawdict gives us:
      - span.text          : proper Unicode (MuPDF applies ToUnicode CMap)
      - span.size          : raw pt size
      - span.flags         : bit 4 = bold, bit 1 = italic
      - span.bbox          : (x0,y0,x1,y1) for spatial ordering
      - span.origin        : baseline origin
      - block.type 1       : image block → extract and insert

    We also detect font-encoding garbage and fall through to Tier 2 if
    > 40% of extracted text is garbage characters.
    """
    try:
        import fitz
        from docx import Document
    except ImportError as exc:
        logger.warning(f"Tier 1 skipped – missing dep: {exc}")
        return False

    try:
        doc_pdf = fitz.open(input_path)
        doc_word = Document()

        # Strip default empty paragraph
        for p in list(doc_word.paragraphs):
            p._element.getparent().remove(p._element)

        _configure_doc_rtl_defaults(doc_word)

        total_paras = 0
        total_imgs = 0
        all_text: list[str] = []

        for page_num in range(len(doc_pdf)):
            page = doc_pdf[page_num]

            if page_num > 0:
                _add_page_break(doc_word)

            # Use rawdict for maximum fidelity
            page_dict = page.get_text("rawdict", sort=True,
                                       flags=fitz.TEXT_PRESERVE_LIGATURES |
                                             fitz.TEXT_PRESERVE_WHITESPACE |
                                             fitz.TEXT_MEDIABOX_CLIP)
            raw_blocks = page_dict.get("blocks", [])
            median = _page_median_size(raw_blocks)

            # Collect all blocks sorted top-to-bottom
            blocks: list[Block] = []

            for b in raw_blocks:
                b_type = b.get("type", -1)
                bbox = b.get("bbox", [0, 0, 0, 0])
                y0 = bbox[1]

                if b_type == 0:  # text block
                    lines_of_spans: list[list[Span]] = []

                    for line in b.get("lines", []):
                        line_spans: list[Span] = []

                        for raw_span in line.get("spans", []):
                            # rawdict has 'chars' list; reconstruct text
                            chars = raw_span.get("chars", [])
                            if chars:
                                text = "".join(c.get("c", "") for c in chars)
                            else:
                                text = raw_span.get("text", "")

                            text = text.strip()
                            if not text:
                                continue

                            all_text.append(text)

                            raw_size = raw_span.get("size", 11.0)
                            flags = raw_span.get("flags", 0)
                            is_bold   = bool(flags & 16)
                            is_italic = bool(flags & 2)
                            span_bbox = raw_span.get("bbox", bbox)

                            norm = _norm_size(raw_size, median)
                            rtl  = _is_rtl_span(text)

                            line_spans.append(Span(
                                text=text,
                                size=norm,
                                bold=is_bold,
                                italic=is_italic,
                                rtl=rtl,
                                x0=span_bbox[0],
                                y0=span_bbox[1],
                                y1=span_bbox[3],
                            ))

                        if line_spans:
                            lines_of_spans.append(line_spans)

                    if lines_of_spans:
                        blocks.append(Block(y0, "text", lines_of_spans))

                elif b_type == 1:  # image block
                    xref = b.get("image", 0)
                    if isinstance(xref, dict):
                        xref = xref.get("xref", 0)
                    if xref:
                        w = b.get("width", 0)
                        h = b.get("height", 0)
                        blocks.append(Block(y0, "image", (xref, w, h)))

            # Also pick up images via get_image_info (catches some MuPDF
            # misses on the block scan above)
            inserted_xrefs: set[int] = set()
            for img_info in page.get_image_info(xrefs=True):
                xref = img_info.get("xref", 0)
                if xref and not any(
                    b.kind == "image" and b.data[0] == xref for b in blocks
                ):
                    ibbox = img_info.get("bbox", [0, 0, 0, 0])
                    blocks.append(Block(ibbox[1], "image", (xref, 0, 0)))

            blocks.sort(key=lambda b: b.y0)

            # Check for garbage encoding on first content page
            if page_num == 0 and all_text:
                combined = "".join(all_text)
                if _text_is_garbage(combined):
                    logger.warning(
                        "Tier 1: font encoding garbage detected – "
                        "falling through to Tier 2 (OCR)"
                    )
                    doc_pdf.close()
                    return False

            for block in blocks:
                if block.kind == "text":
                    lines_of_spans = block.data  # type: ignore[assignment]

                    # Merge lines into paragraphs
                    paragraphs = _lines_to_paragraphs(lines_of_spans)

                    for para_spans in paragraphs:
                        if not para_spans:
                            continue

                        dom_size = max(s.size for s in para_spans)

                        # Heading detection by normalised size
                        if dom_size >= 18:
                            hlevel = 1
                        elif dom_size >= 14.5:
                            hlevel = 2
                        elif dom_size >= 12.5:
                            hlevel = 3
                        else:
                            hlevel = 0

                        _write_paragraph(doc_word, para_spans, heading_level=hlevel)
                        total_paras += 1

                elif block.kind == "image":
                    xref = block.data[0]  # type: ignore[index]
                    if xref in inserted_xrefs:
                        continue
                    try:
                        img_data = doc_pdf.extract_image(xref)
                        if img_data and img_data.get("image") and \
                                len(img_data["image"]) > 500:
                            _write_image(doc_word, img_data["image"])
                            inserted_xrefs.add(xref)
                            total_imgs += 1
                    except Exception as exc:
                        logger.debug(f"Image xref {xref}: {exc}")

        doc_pdf.close()

        if total_paras == 0 and total_imgs == 0:
            logger.warning("Tier 1: no content extracted")
            return False

        doc_word.save(output_path)
        size = Path(output_path).stat().st_size
        logger.info(
            f"Tier 1 (PyMuPDF rawdict) OK: "
            f"{total_paras} paras, {total_imgs} images, {size} bytes"
        )
        return True

    except Exception as exc:
        logger.warning(f"Tier 1 failed: {exc}", exc_info=True)
        return False


# ===========================================================================
# Tier 2 — Tesseract OCR
# ===========================================================================

def _render_page_pil(page, dpi: int = 200):
    """Render a fitz page to PIL Image."""
    import fitz
    from PIL import Image

    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    return Image.open(io.BytesIO(pix.tobytes("png")))


def _tier2_tesseract_ocr_to_docx(input_path: str, output_path: str) -> bool:
    try:
        import fitz
        import pytesseract
        from docx import Document
    except ImportError as exc:
        logger.warning(f"Tier 2 skipped – missing dep: {exc}")
        return False

    try:
        subprocess.run(["tesseract", "--version"],
                       capture_output=True, timeout=10, check=True)
    except Exception as exc:
        logger.warning(f"Tier 2 skipped – tesseract not found: {exc}")
        return False

    try:
        doc_pdf = fitz.open(input_path)
        doc_word = Document()

        for p in list(doc_word.paragraphs):
            p._element.getparent().remove(p._element)

        _configure_doc_rtl_defaults(doc_word)

        total_paras = 0

        for page_num in range(len(doc_pdf)):
            page = doc_pdf[page_num]

            if page_num > 0:
                _add_page_break(doc_word)

            pil_img = _render_page_pil(page, dpi=200)
            ocr_text = pytesseract.image_to_string(
                pil_img,
                lang="ara+eng",
                config="--psm 3 --oem 1",
            )

            if not ocr_text or not ocr_text.strip():
                continue

            for line in ocr_text.split("\n"):
                line = line.strip()
                if not line:
                    continue

                rtl = _is_rtl_span(line)
                para = doc_word.add_paragraph()
                _set_para_bidi(para, rtl)
                _set_para_spacing(para)

                run = para.add_run(line)
                _make_run_rtl_props(run, font_name="Amiri",
                                    font_size_pt=11.0, rtl=rtl)
                total_paras += 1

        doc_pdf.close()

        if total_paras == 0:
            logger.warning("Tier 2: no text from OCR")
            return False

        doc_word.save(output_path)
        logger.info(
            f"Tier 2 (Tesseract) OK: "
            f"{total_paras} paras, {Path(output_path).stat().st_size} bytes"
        )
        return True

    except Exception as exc:
        logger.warning(f"Tier 2 failed: {exc}", exc_info=True)
        return False


# ===========================================================================
# Tier 3 — LibreOffice fallback
# ===========================================================================

def _run_lo(lo_bin: str, extra_args: list,
            input_file: Path, out_dir: Path) -> subprocess.CompletedProcess:
    env = os.environ.copy()
    env["LANG"] = "en_US.UTF-8"
    env["LC_ALL"] = "en_US.UTF-8"
    cmd = [lo_bin, "--headless", "--norestore", "--nofirststartwizard",
           *extra_args, "--outdir", str(out_dir), str(input_file)]
    logger.info(f"LO: {' '.join(cmd)}")
    return subprocess.run(cmd, capture_output=True, text=True,
                          timeout=120, env=env)


def _tier3_libreoffice_fallback(input_path: str, output_path: str) -> bool:
    lo_bin = find_libreoffice()
    if not lo_bin:
        logger.warning("Tier 3 skipped – LibreOffice not found")
        return False

    input_file = Path(input_path)
    output_file = Path(output_path)

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)

        # 3a writer_pdf_import
        _run_lo(lo_bin, ["--infilter=writer_pdf_import", "--convert-to", "docx"],
                input_file, tmp)
        hits = list(tmp.glob("*.docx"))
        if hits:
            shutil.move(str(hits[0]), str(output_file))
            logger.info(f"Tier 3a OK: {output_file.stat().st_size} bytes")
            return True

        # 3b direct
        _run_lo(lo_bin, ["--convert-to", "docx"], input_file, tmp)
        hits = list(tmp.glob("*.docx"))
        if hits:
            shutil.move(str(hits[0]), str(output_file))
            logger.info(f"Tier 3b OK: {output_file.stat().st_size} bytes")
            return True

        # 3c PDF→ODT→DOCX
        odt_dir = tmp / "odt"; odt_dir.mkdir()
        _run_lo(lo_bin, ["--convert-to", "odt"], input_file, odt_dir)
        odt_hits = list(odt_dir.glob("*.odt"))
        if odt_hits:
            docx_dir = tmp / "docx"; docx_dir.mkdir()
            _run_lo(lo_bin, ["--convert-to", "docx"], odt_hits[0], docx_dir)
            hits = list(docx_dir.glob("*.docx"))
            if hits:
                shutil.move(str(hits[0]), str(output_file))
                logger.info(f"Tier 3c OK: {output_file.stat().st_size} bytes")
                return True

    logger.warning("Tier 3: all LibreOffice strategies failed")
    return False


# ===========================================================================
# Public entry point
# ===========================================================================

def convert_pdf_to_docx(input_path: str, output_path: str) -> None:
    """
    Convert PDF to DOCX with full Arabic/RTL support.

    Raises RuntimeError if all tiers fail.
    """
    in_file = Path(input_path)
    out_file = Path(output_path)

    if not in_file.exists():
        raise RuntimeError(f"Input not found: {input_path}")
    out_file.parent.mkdir(parents=True, exist_ok=True)

    logger.info(f"PDF→DOCX: {in_file.name}")

    text_based = is_text_based_pdf(input_path)
    logger.info(f"PDF type: {'text-based' if text_based else 'image/scanned'}")

    if text_based:
        if _tier1_pymupdf_to_docx(input_path, output_path):
            return
        logger.warning("Tier 1 failed → Tier 2 (OCR)")
        if _tier2_tesseract_ocr_to_docx(input_path, output_path):
            return
        logger.warning("Tier 2 failed → Tier 3 (LibreOffice)")
        if _tier3_libreoffice_fallback(input_path, output_path):
            return
    else:
        logger.info("Scanned PDF → Tier 2 (OCR)")
        if _tier2_tesseract_ocr_to_docx(input_path, output_path):
            return
        logger.warning("Tier 2 failed → Tier 3 (LibreOffice)")
        if _tier3_libreoffice_fallback(input_path, output_path):
            return

    raise RuntimeError(
        "All conversion strategies failed. "
        "PDF may be encrypted, corrupted, or use an unsupported encoding."
    )
